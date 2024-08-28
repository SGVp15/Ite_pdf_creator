import os.path

from docx import Document

from contact import Contact
from config import TEMPLATES_DIR, OUT_DOCX_DIR, OUT_PDF_DIR


def replace_docx_text(document, old_text: str, new_text: str):
    old_text = str(old_text)
    new_text = str(new_text)
    old_text = '{' + old_text + '}'
    section = document.sections[0]
    header = section.header
    footer = section.footer
    for doc_obj in (document, header, footer):
        docx_replace_regex(doc_obj, old_text, new_text)


def docx_replace_regex(doc_obj, old_text: str, new_text: str):
    old_text = str(old_text)
    new_text = str(new_text)
    # Замена в параграфах
    for paragraph in doc_obj.paragraphs:
        if old_text not in paragraph.text:
            continue
        runs = paragraph.runs
        if len(runs) == 0:
            continue
        if len(runs) > 1:
            text = ''
            for run in runs:
                text += run.text
            runs[0].text = text
            for i in range(1, len(runs)):
                runs[i].text = ''
        runs[0].text = runs[0].text.replace(old_text, new_text)

    # Замена в таблицах
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, old_text, new_text)


def create_docx(contact: Contact):
    for docx_template in contact.docx_list_files_name_templates:
        document = Document(docx=str(os.path.join(TEMPLATES_DIR, docx_template)))
        # Gender Пол
        gender_text = 'прошел'
        if contact.gender.lower() in ('ж', 'f'):
            gender_text = 'прошла'

        replaces_dict = {
            'Gender': gender_text,
            'NameRus': contact.name_rus,
            'NameEng': contact.name_eng,
            'Number': contact.sert_number,
            'CourseRus': contact.course.name_rus,
            'CourseEng': contact.course.name_eng,
            'HoursRus': contact.course.hour_rus,
            'HoursEng': contact.course.hour_eng,
            'CourseDateRus': contact.course_date_rus,
            'CourseDateEng': contact.course_date_eng,
            'IssueDateRus': contact.issue_date_rus,
            'Year': contact.year}

        # Замена всех полей
        for k, v in replaces_dict.items():
            replace_docx_text(document, old_text=k, new_text=v)

        path = contact.files_out_docx.get(docx_template)
        contact.create_dirs()

        document.save(path)
